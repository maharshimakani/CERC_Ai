"""
Output Generator Module
Generates final CSR outputs in various formats (TXT, DOCX, PDF).
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Dict, List, Optional, TYPE_CHECKING, Any
from datetime import datetime

if TYPE_CHECKING:
    from docx import Document

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    FPDF = None
    PDF_AVAILABLE = False


class OutputGenerator:
    """Generates CSR output files in various formats."""
    
    # Section order for the final document (ICH E3 order)
    SECTION_ORDER = [
        "synopsis",               # Synopsis
        "introduction",           # Section 1
        "ethics",                 # Section 4
        "study_objectives",       # Section 5
        "investigators_sites",    # Section 6
        "study_design",           # Section 9.1
        "inclusion_exclusion",    # Section 9.3
        "treatments",             # Section 9.4
        "endpoints",              # Section 9.4.1
        "study_population",       # Section 10.1
        "demographics",           # Section 10.1.4
        "efficacy_evaluation",    # Section 10
        "statistical_methods",    # Section 11
        "safety_evaluation",      # Section 12
        "adverse_events",         # Section 12.2
        "discussion_conclusions"  # Section 13
    ]
    
    # Section display names
    SECTION_NAMES = {
        "synopsis": "Synopsis",
        "introduction": "1. Introduction",
        "ethics": "4. Ethics",
        "study_objectives": "5. Study Objectives",
        "investigators_sites": "6. Investigators and Study Sites",
        "study_design": "9.1 Study Design",
        "inclusion_exclusion": "9.3 Selection of Study Population",
        "treatments": "9.4 Study Treatments",
        "endpoints": "9.4.1 Efficacy and Safety Variables",
        "study_population": "10.1 Subject Disposition",
        "demographics": "10.1.4 Demographics and Baseline Characteristics",
        "efficacy_evaluation": "10. Efficacy Evaluation",
        "statistical_methods": "11. Statistical Methods",
        "safety_evaluation": "12. Safety Evaluation",
        "adverse_events": "12.2 Adverse Events",
        "discussion_conclusions": "13. Discussion and Overall Conclusions"
    }
    
    def __init__(self, output_dir: Path):
        """
        Initialize output generator.
        
        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.sections_dir = self.output_dir / "csr_sections"
        self.sections_dir.mkdir(parents=True, exist_ok=True)
        
    @staticmethod
    def sanitize_markdown(text: str) -> str:
        """
        Strip all Markdown syntax from CSR content.
        This is a safety net to ensure no markup ever reaches final output.
        
        Removes: # headings, ** bold, * italics, - bullets, ``` code blocks,
        > blockquotes, horizontal rules (---), and other markup artifacts.
        """
        if not text:
            return text
        
        # Remove triple-backtick code fences entirely
        text = re.sub(r'```[\w]*\n?', '', text)
        
        lines = text.split('\n')
        cleaned = []
        
        for line in lines:
            # Strip markdown heading prefixes (####, ###, ##, #)
            line = re.sub(r'^#{1,6}\s+', '', line)
            # Strip mid-line heading marks (e.g., "Some text ### Heading")
            line = re.sub(r'\s*#{1,6}\s+', ' ', line)
            # Strip blockquote markers
            line = re.sub(r'^>\s*', '', line)
            # Strip bold markup **text** -> text
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Strip italic markup *text* -> text (but not bullet lines)
            if not line.strip().startswith('* '):
                line = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', line)
            # Convert bullet lines (- text or * text) to plain text
            line = re.sub(r'^\s*[-*]\s+', '', line)
            # Strip inline code backticks
            line = re.sub(r'`([^`]*)`', r'\1', line)
            # Strip horizontal rules
            if re.match(r'^\s*[-*_]{3,}\s*$', line):
                line = ''
            
            cleaned.append(line)
        
        return '\n'.join(cleaned)
    
    @staticmethod
    def inject_figure_placeholders(text: str, source_filename: str = "") -> str:
        """
        Detect references to figures and tables in generated text and insert
        placeholder text indicating the reader should refer to the source document.
        
        Args:
            text: Generated section text
            source_filename: Name of the source document for reference
            
        Returns:
            Text with figure/table placeholders injected
        """
        if not text:
            return text
        
        source_ref = f" in {source_filename}" if source_filename else ""
        
        # Detect "Figure X", "Table X", "Chart X", "Graph X" references
        def _replace_ref(match):
            ref_type = match.group(1)  # Figure, Table, etc.
            ref_num = match.group(2)   # The number
            return f"{ref_type} {ref_num} [See {ref_type} {ref_num}{source_ref}]"
        
        text = re.sub(
            r'\b(Figure|Table|Chart|Graph|Diagram)\s+(\d+(?:\.\d+)?)\b',
            _replace_ref,
            text,
            flags=re.IGNORECASE
        )
        
        return text
    
    @staticmethod
    def validate_no_markdown(text: str) -> List[str]:
        """
        Check for any remaining Markdown syntax in final output.
        Returns list of violations found.
        """
        violations = []
        if not text:
            return violations
            
        for i, line in enumerate(text.split('\n'), 1):
            if re.match(r'^#{1,6}\s+', line):
                violations.append(f"Line {i}: Markdown heading detected: {line[:50]}")
            if '**' in line:
                violations.append(f"Line {i}: Bold markup detected: {line[:50]}")
        
        return violations

    def save_section_txt(
        self,
        section_id: str,
        content: str,
        metadata: Dict = None
    ) -> Path:
        """
        Save a single section as TXT file.
        
        Args:
            section_id: The CSR section ID
            content: The section content
            metadata: Optional metadata to include
            
        Returns:
            Path to saved file
        """
        # Sanitize content before saving
        content = self.sanitize_markdown(content)
        
        output_file = self.sections_dir / f"{section_id}.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            # Write header
            f.write(f"{'='*60}\n")
            f.write(f"CSR Section: {self.SECTION_NAMES.get(section_id, section_id)}\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            if metadata:
                sources = metadata.get("source_documents", [])
                if sources:
                    f.write(f"Source Documents: {', '.join(sources)}\n")
            f.write(f"{'='*60}\n\n")
            
            # Write content
            f.write(content)
            
        print(f"Saved: {output_file.name}")
        return output_file
    
    def save_all_sections_txt(
        self,
        generated_sections: Dict[str, Dict]
    ) -> List[Path]:
        """
        Save all generated sections as individual TXT files.
        
        Args:
            generated_sections: Dictionary of generated section data
            
        Returns:
            List of saved file paths
        """
        saved_files = []
        
        for section_id in self.SECTION_ORDER:
            if section_id in generated_sections:
                section_data = generated_sections[section_id]
                file_path = self.save_section_txt(
                    section_id,
                    section_data.get("final_text", ""),
                    section_data
                )
                saved_files.append(file_path)
                
        return saved_files
    
    def generate_docx(
        self,
        generated_sections: Dict[str, Dict],
        study_title: str = "Clinical Study Report",
        protocol_number: str = "PROTOCOL-XXX"
    ) -> Path:
        """
        Generate a complete CSR DOCX document.
        
        Args:
            generated_sections: Dictionary of generated section data
            study_title: Title of the study
            protocol_number: Protocol number
            
        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
            
        doc = DocxDocument()
        
        # Title page
        title = doc.add_heading(study_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        subtitle = doc.add_paragraph("CLINICAL STUDY REPORT")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        protocol = doc.add_paragraph(f"Protocol Number: {protocol_number}")
        protocol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        date = doc.add_paragraph(f"Report Date: {datetime.now().strftime('%d %B %Y')}")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        note = doc.add_paragraph("This report was generated using AI-assisted writing tools.")
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.runs[0].italic = True
        
        # Page break before content
        doc.add_page_break()
        
        # Generate actual Table of Contents
        doc.add_heading("TABLE OF CONTENTS", level=1)
        toc_entries = self._generate_toc_entries(generated_sections)
        for entry in toc_entries:
            toc_para = doc.add_paragraph(entry)
            toc_para.paragraph_format.left_indent = Inches(0.5)
        
        # Add document inventory if available
        if hasattr(self, '_document_inventory') and self._document_inventory:
            doc.add_paragraph()
            doc.add_heading("Source Documents", level=2)
            inv = self._document_inventory
            
            if inv.get("provided"):
                doc.add_paragraph("Documents provided and processed:")
                for doc_name in inv["provided"]:
                    p = doc.add_paragraph(f"• {doc_name}", style='List Bullet')
                    
            if inv.get("expected_but_missing"):
                doc.add_paragraph("Expected documents not found:")
                for doc_type in inv["expected_but_missing"]:
                    p = doc.add_paragraph(f"• {doc_type} (not provided)", style='List Bullet')
                    
            if inv.get("extraction_issues"):
                doc.add_paragraph("Documents with extraction issues:")
                for issue in inv["extraction_issues"]:
                    p = doc.add_paragraph(f"• {issue}", style='List Bullet')
        
        doc.add_page_break()
        
        # Add each section
        for section_id in self.SECTION_ORDER:
            if section_id in generated_sections:
                section_data = generated_sections[section_id]
                content = section_data.get("final_text", "")
                
                if content and "[No content available" not in content:
                    # Sanitize markdown before adding to document
                    content = self.sanitize_markdown(content)
                    # Parse and add content
                    self._add_section_to_doc(doc, content)
                    doc.add_paragraph()  # Spacing between sections
                    
        # Save document
        output_file = self.output_dir / "csr_draft.docx"
        doc.save(str(output_file))
        
        print(f"\nGenerated DOCX: {output_file}")
        return output_file
    
    def _generate_toc_entries(self, generated_sections: Dict) -> List[str]:
        """
        Generate table of contents entries from generated sections.
        
        Args:
            generated_sections: Dictionary of generated sections
            
        Returns:
            List of TOC entry strings
        """
        entries = []
        for section_id in self.SECTION_ORDER:
            if section_id in generated_sections:
                section_name = self.SECTION_NAMES.get(section_id, section_id)
                entries.append(section_name)
        return entries
    
    def set_document_inventory(self, inventory: Dict) -> None:
        """
        Set the document inventory for inclusion in the report.
        
        Args:
            inventory: Dictionary with 'provided', 'expected_but_missing', 'extraction_issues'
        """
        self._document_inventory = inventory
    
    def _is_section_heading(self, line: str) -> bool:
        """
        Detect if a line is a section heading (plain numbered title).
        
        Examples:
            '5 STUDY OBJECTIVES' -> True
            '9.1 Study Design' -> True
            '9.1.1 Overall Study Design' -> True
            '11.3 Statistical Methods for Primary Endpoint' -> True
            '1. First criterion' -> False (numbered list item)
        """
        # Match patterns like: 5, 9.1, 9.1.1, 11.3, etc. followed by a title
        match = re.match(r'^(\d+(?:\.\d+)*)\s+([A-Z])', line)
        if match:
            # Ensure it's not a numbered list like "1. First item"
            num_part = match.group(1)
            # Single digit followed by period+space is a list item, not a heading
            if re.match(r'^\d+$', num_part) and not line[len(num_part):].strip()[0].isupper():
                return False
            return True
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """
        Determine DOCX heading level from section number depth.
        
        '5 OBJECTIVES' -> level 1
        '9.1 Study Design' -> level 2
        '9.1.1 Overall Design' -> level 3
        '11.3.2 Subanalysis' -> level 3
        """
        match = re.match(r'^(\d+(?:\.\d+)*)\s+', line)
        if match:
            parts = match.group(1).split('.')
            depth = len(parts)
            return min(depth, 4)  # Cap at level 4
        return 2  # Default
    
    def _add_section_to_doc(self, doc: 'Document', content: str) -> None:
        """
        Add section content to DOCX using plain numbered headings.
        
        Args:
            doc: The Document object
            content: Section content with plain numbered headings
        """
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Strip any residual markdown headings (safety net)
            line = re.sub(r'^#{1,6}\s+', '', line)
            # Strip bold markup
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                
            # Numbered section headings (e.g., "9.1.1 Title")
            if self._is_section_heading(line):
                level = self._get_heading_level(line)
                doc.add_heading(line, level=level)
            # Numbered list items (e.g., "1. First criterion")
            elif re.match(r'^\d+\.\s+', line):
                doc.add_paragraph(line, style='List Number')
            # Regular paragraph
            else:
                doc.add_paragraph(line)
    
    def _init_pdf_fonts(self, pdf: 'FPDF') -> str:
        """
        Initialize a Unicode-capable font for the PDF.
        Tries Windows Arial first, then falls back to Helvetica with char replacement.
        
        Returns:
            Font family name to use throughout the document.
        """
        # Try Windows Arial (supports Unicode clinical symbols)
        arial_paths = [
            Path(r'C:\Windows\Fonts\arial.ttf'),
            Path(r'C:\Windows\Fonts\Arial.ttf'),
        ]
        for font_path in arial_paths:
            if font_path.exists():
                try:
                    pdf.add_font('Arial', '', str(font_path), uni=True)
                    # Also add bold and italic variants
                    bold_path = font_path.parent / 'arialbd.ttf'
                    italic_path = font_path.parent / 'ariali.ttf'
                    if bold_path.exists():
                        pdf.add_font('Arial', 'B', str(bold_path), uni=True)
                    if italic_path.exists():
                        pdf.add_font('Arial', 'I', str(italic_path), uni=True)
                    return 'Arial'
                except Exception:
                    pass
        
        # Fallback: use Helvetica but sanitize Unicode chars
        self._pdf_needs_sanitize = True
        return 'Helvetica'
    
    @staticmethod
    def _sanitize_for_ascii_font(text: str) -> str:
        """
        Replace common Unicode clinical symbols with ASCII equivalents
        for use with non-Unicode fonts.
        """
        replacements = {
            '\u2265': '>=',   # ≥
            '\u2264': '<=',   # ≤  
            '\u00b1': '+/-',  # ±
            '\u00b5': 'u',    # µ (micro)
            '\u2013': '-',    # en dash
            '\u2014': '--',   # em dash
            '\u2018': "'",    # left single quote
            '\u2019': "'",    # right single quote
            '\u201c': '"',    # left double quote
            '\u201d': '"',    # right double quote
            '\u2022': '-',    # bullet
            '\u03b1': 'alpha',  # α
            '\u03b2': 'beta',   # β
            '\u00d7': 'x',    # ×
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        return text
    
    def _pdf_text(self, text: str) -> str:
        """Prepare text for PDF rendering, sanitizing if needed."""
        if getattr(self, '_pdf_needs_sanitize', False):
            return self._sanitize_for_ascii_font(text)
        return text
    
    def generate_pdf(
        self,
        generated_sections: Dict[str, Dict],
        study_title: str = "Clinical Study Report",
        protocol_number: str = "PROTOCOL-XXX"
    ) -> Path:
        """
        Generate a complete CSR PDF document.
        Content is identical to the DOCX output.
        Uses Unicode fonts to support clinical notation (>=, <=, +/-, etc.).
        
        Args:
            generated_sections: Dictionary of generated section data
            study_title: Title of the study
            protocol_number: Protocol number
            
        Returns:
            Path to generated PDF file
        """
        if not PDF_AVAILABLE:
            raise ImportError("fpdf2 is not installed. Run: pip install fpdf2")
        
        self._pdf_needs_sanitize = False
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=25)
        
        # Initialize Unicode font
        font = self._init_pdf_fonts(pdf)
        self._pdf_font = font
        
        # Title page
        pdf.add_page()
        pdf.set_font(font, 'B', 20)
        pdf.cell(0, 60, '', ln=True)  # Spacing
        pdf.cell(0, 12, self._pdf_text(study_title), ln=True, align='C')
        pdf.ln(8)
        pdf.set_font(font, 'B', 16)
        pdf.cell(0, 10, 'CLINICAL STUDY REPORT', ln=True, align='C')
        pdf.ln(8)
        pdf.set_font(font, '', 12)
        pdf.cell(0, 8, f'Protocol Number: {protocol_number}', ln=True, align='C')
        pdf.ln(4)
        pdf.cell(0, 8, f"Report Date: {datetime.now().strftime('%d %B %Y')}", ln=True, align='C')
        pdf.ln(8)
        pdf.set_font(font, 'I', 10)
        pdf.cell(0, 8, 'This report was generated using AI-assisted writing tools.', ln=True, align='C')
        
        # Table of Contents page
        pdf.add_page()
        pdf.set_font(font, 'B', 16)
        pdf.cell(0, 12, 'TABLE OF CONTENTS', ln=True)
        pdf.ln(6)
        pdf.set_font(font, '', 11)
        toc_entries = self._generate_toc_entries(generated_sections)
        for entry in toc_entries:
            pdf.cell(10, 8, '', ln=False)  # Indent
            pdf.cell(0, 8, entry, ln=True)
        
        # Content pages
        for section_id in self.SECTION_ORDER:
            if section_id in generated_sections:
                section_data = generated_sections[section_id]
                content = section_data.get("final_text", "")
                
                if content and "[No content available" not in content:
                    content = self.sanitize_markdown(content)
                    pdf.add_page()
                    self._add_section_to_pdf(pdf, content)
        
        # Save PDF
        output_file = self.output_dir / "csr_draft.pdf"
        pdf.output(str(output_file))
        
        print(f"\nGenerated PDF: {output_file}")
        return output_file
    
    def _add_section_to_pdf(self, pdf: 'FPDF', content: str) -> None:
        """
        Add section content to PDF with proper heading formatting.
        
        Args:
            pdf: The FPDF object
            content: Section content with plain numbered headings
        """
        font = getattr(self, '_pdf_font', 'Helvetica')
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue
            
            # Strip residual markdown (safety net)
            line = re.sub(r'^#{1,6}\s+', '', line)
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Prepare for font rendering
            line = self._pdf_text(line)
            
            if self._is_section_heading(line):
                level = self._get_heading_level(line)
                if level == 1:
                    pdf.set_font(font, 'B', 16)
                    pdf.ln(6)
                elif level == 2:
                    pdf.set_font(font, 'B', 13)
                    pdf.ln(4)
                else:
                    pdf.set_font(font, 'B', 11)
                    pdf.ln(3)
                pdf.multi_cell(0, 7, line)
                pdf.ln(2)
            elif re.match(r'^\d+\.\s+', line):
                # Numbered list item
                pdf.set_font(font, '', 10)
                pdf.cell(8, 6, '', ln=False)  # Indent
                pdf.multi_cell(0, 6, line)
                pdf.ln(1)
            else:
                # Regular paragraph
                pdf.set_font(font, '', 10)
                pdf.multi_cell(0, 6, line)
                pdf.ln(2)
    
    def generate_traceability_log(
        self,
        generated_sections: Dict[str, Dict],
        validation_results: Dict = None
    ) -> Path:
        """
        Generate a traceability log documenting the generation process.
        
        Args:
            generated_sections: Dictionary of generated section data
            validation_results: Optional validation results
            
        Returns:
            Path to log file
        """
        log_data = {
            "generation_timestamp": datetime.now().isoformat(),
            "sections_generated": len(generated_sections),
            "sections": {}
        }
        
        for section_id, section_data in generated_sections.items():
            log_data["sections"][section_id] = {
                "section_name": self.SECTION_NAMES.get(section_id, section_id),
                "source_documents": section_data.get("source_documents", []),
                "output_length": len(section_data.get("final_text", "")),
                "has_content": "[No content available" not in section_data.get("final_text", "")
            }
            
        if validation_results:
            log_data["validation"] = validation_results
            
        output_file = self.output_dir / "generation_log.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(log_data, f, indent=2)
            
        print(f"✓ Saved: {output_file.name}")
        return output_file
    
    def generate_summary_report(
        self,
        generated_sections: Dict[str, Dict],
        validation_results: Dict = None
    ) -> Path:
        """
        Generate a human-readable summary report.
        
        Args:
            generated_sections: Dictionary of generated section data
            validation_results: Optional validation results
            
        Returns:
            Path to summary file
        """
        output_file = self.output_dir / "generation_summary.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("="*60 + "\n")
            f.write("CSR GENERATION SUMMARY REPORT\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*60 + "\n\n")
            
            f.write("SECTIONS GENERATED:\n")
            f.write("-"*40 + "\n")
            
            for section_id in self.SECTION_ORDER:
                if section_id in generated_sections:
                    section_data = generated_sections[section_id]
                    name = self.SECTION_NAMES.get(section_id, section_id)
                    length = len(section_data.get("final_text", ""))
                    sources = section_data.get("source_documents", [])
                    
                    status = "✓" if "[No content available" not in section_data.get("final_text", "") else "⚠"
                    
                    f.write(f"\n{status} {name}\n")
                    f.write(f"   Length: {length} characters\n")
                    f.write(f"   Sources: {', '.join(sources) if sources else 'None'}\n")
                    
            if validation_results:
                f.write("\n\nVALIDATION SUMMARY:\n")
                f.write("-"*40 + "\n")
                f.write(f"Sections validated: {validation_results.get('total_sections_validated', 0)}\n")
                f.write(f"Sections passed: {validation_results.get('sections_passed', 0)}\n")
                f.write(f"Total errors: {validation_results.get('total_errors', 0)}\n")
                f.write(f"Total warnings: {validation_results.get('total_warnings', 0)}\n")
                
            f.write("\n" + "="*60 + "\n")
            f.write("END OF REPORT\n")
            
        print(f"✓ Saved: {output_file.name}")
        return output_file
