"""
Document Loader Module
Handles loading of PDF and DOCX files from the input directory.
Includes OCR fallback for scanned PDFs and table extraction.
"""

import os
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import json

logger = logging.getLogger(__name__)

# Try importing document processing libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Optional: pdfplumber for table extraction
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Optional: pytesseract for OCR
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


class DocumentLoader:
    """Loads clinical documents from the input directory."""
    
    # Document type detection keywords (expanded)
    DOCUMENT_TYPE_PATTERNS = {
        "protocol": [
            "clinical study protocol",
            "study protocol",
            "protocol",
            "investigational plan",
            "clinical trial protocol"
        ],
        "sap": [
            "statistical analysis plan",
            "analysis plan",
            "sap",
            "statistical methods"
        ],
        "csr": [
            "clinical study report",
            "study report",
            "final study report",
            "integrated summary"
        ],
        "ib": [
            "investigator brochure",
            "investigator's brochure",
            "ib"
        ],
        "monitoring_plan": [
            "monitoring plan",
            "clinical monitoring",
            "site monitoring"
        ],
        "cec_charter": [
            "clinical events committee",
            "cec charter",
            "endpoint adjudication",
            "adjudication charter"
        ],
        "icf": [
            "informed consent",
            "consent form",
            "patient information"
        ],
        "tlf": [
            "tables listings figures",
            "tables and listings",
            "tabulation",
            "data listings"
        ]
    }
    
    def __init__(self, input_dir: Path):
        """
        Initialize the document loader.
        
        Args:
            input_dir: Path to the input documents directory
        """
        self.input_dir = Path(input_dir)
        self.loaded_documents: Dict[str, Dict] = {}

    def load_user_documents(self, paths: List[Path]) -> List[Dict[str, Any]]:
        """
        Load ONLY the provided user-uploaded files.

        This enforces separation between:
          - static knowledge layer (`resources/`) handled elsewhere (KnowledgeEngine)
          - dynamic user evidence handled by this loader

        Returns objects shaped like:
        {
          filename,
          filepath,
          document_type,
          raw_text,
          pages: [{"page_num", "text"}],
          tables: [{"caption", "rows"}]
        }
        """
        user_docs: List[Dict[str, Any]] = []
        for p in paths:
            file_path = Path(p)
            if not file_path.exists() or not file_path.is_file():
                continue

            loaded = self.load_document(file_path)

            raw_text = loaded.get("full_text", "") or ""
            doc_type = loaded.get("document_type", "unknown")
            ext = loaded.get("extension", "")

            pages: List[Dict[str, Any]] = []
            structure = loaded.get("structure", [])
            if ext == ".pdf" and isinstance(structure, list):
                for pg in structure:
                    pages.append(
                        {
                            "page_num": int(pg.get("page_number", 1) or 1),
                            "text": str(pg.get("text", "") or ""),
                        }
                    )
            else:
                # For DOCX (and DOC), we don't have true page boundaries;
                # represent the whole document as a single page for provenance.
                pages = [{"page_num": 1, "text": raw_text}]

            user_docs.append(
                {
                    "filename": str(loaded.get("filename", file_path.name)),
                    "filepath": str(loaded.get("filepath", str(file_path))),
                    "document_type": doc_type,
                    "raw_text": raw_text,
                    # Compatibility fields for existing pipeline components.
                    "full_text": raw_text,
                    "char_count": len(raw_text),
                    "word_count": len(raw_text.split()),
                    "pages": pages,
                    # Table extraction is optional; current pipeline captures tables in raw text.
                    "tables": [],
                    # Keep legacy 'structure' field (used by some components as provenance).
                    "structure": pages if pages else [],
                }
            )

        return user_docs

    def _classify_resource(self, filename: str, text: str) -> str:
        """Classify static knowledge document type by filename and text."""
        f = filename.lower()
        t = text[:5000].lower()
        combined = f + " " + t
        
        if "clinical investigation report" in combined or "clinical study report" in combined or "final report" in combined:
            return "clinical_report_example"
        elif "statistical analysis plan" in combined or "statistical report" in combined or "sap" in f:
            return "statistical_reference"
        elif "clinical investigation plan" in combined or "protocol" in combined or "cip" in f:
            return "protocol_reference"
        elif "template" in combined or "format" in combined or "section" in combined:
            return "template_reference"
        else:
            return "general_reference"

    def load_resources(self, resources_path: Path) -> List[Dict]:
        """
        Load static knowledge documents from resources folder.
        These are NOT evidence documents.
        """
        resources = []
        if not resources_path.exists():
            return resources
            
        for file in resources_path.glob("*.*"):
            if file.suffix.lower() not in [".pdf", ".docx", ".doc"]:
                continue
                
            doc = self.load_document(file)
            doc["is_resource"] = True
            doc["doc_type"] = "resource"
            text = doc.get("full_text", "")
            doc["resource_type"] = self._classify_resource(file.name, text)
            conf = doc.get("extraction_confidence", "UNKNOWN")
            resources.append(doc)
            
            print(f"[RESOURCE LOAD] {file.name} | type={doc['resource_type']} | len={len(text)}")
            print(f"[RESOURCE LOAD] Extraction confidence: {conf}")
            
        return resources

    # ── Legacy logic below (preserved for back-compat) ───────────────
    def get_available_files(self) -> List[Path]:
        """Get list of all supported files in the input directory and subdirectories."""
        supported_extensions = ['.pdf', '.docx', '.doc']
        files = []
        
        if not self.input_dir.exists():
            return files
            
        for ext in supported_extensions:
            # Recursive glob to pick up subdirectories (e.g. Appendices/)
            files.extend(self.input_dir.rglob(f"*{ext}"))
            
        return sorted(files)
    
    def load_pdf(self, file_path: Path) -> Tuple[str, List[Dict]]:
        """
        Load a PDF file and extract text with structure.

        Extraction strategy (ordered fallback chain):
          1. pdfplumber  — best for complex layouts/tables (PRIMARY)
          2. PyMuPDF     — fast, robust for standard PDFs (SECONDARY)
          3. Raw bytes   — genuine text recovery only (alpha ratio check required)

        Each method is tried in order. Fallback is triggered when the
        extracted text length is < 100 chars.

        If ALL methods fail, full_text remains empty ("").
        Failure is signalled through metadata only:
          extraction_failed = True
          extraction_method = "none"
          quality_tag      = "FAILED"

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (full_text, pages_data)  — full_text may be empty string
        """
        full_text = ""
        pages_data = []
        _extraction_method = "none"  # tracks which method succeeded
        baseline_len = 100  # Default non-zero fallback

        # ── Get Baseline via PyMuPDF ───────────────────────────────────
        try:
            import fitz
            doc = fitz.open(str(file_path))
            baseline_len = max(100, sum(len(page.get_text()) for page in doc))
            doc.close()
        except:
            baseline_len = max(100, file_path.stat().st_size / 3)

        # ── Strategy 1: pdfplumber (PRIMARY) ───────────────────────────
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                page_texts = [page.extract_text() or "" for page in pdf.pages]
                candidate = "\n".join(page_texts).strip()
                if len(candidate) >= 100:
                    full_text = candidate
                    _extraction_method = "pdfplumber"
                    for i, t in enumerate(page_texts, 1):
                        pages_data.append({
                            "page_number": i,
                            "text": t,
                            "char_count": len(t),
                            "ocr_applied": False,
                        })
                    logger.info("[EXTRACT] pdfplumber succeeded: %d chars", len(full_text))
                else:
                    logger.info(
                        "[EXTRACT] pdfplumber weak (%d chars) — falling back to PyMuPDF",
                        len(candidate),
                    )
        except Exception as _e:
            logger.warning("[EXTRACT] pdfplumber failed: %s", _e)

        # ── Strategy 2: PyMuPDF (SECONDARY) ────────────────────────────
        if len(full_text) < 100:
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(file_path))
                page_texts = [page.get_text() for page in doc]
                candidate = "\n".join(page_texts).strip()
                if len(candidate) >= 100:
                    full_text = candidate
                    _extraction_method = "pymupdf"
                    pages_data = []  # reset to avoid mixing pages from two methods
                    for i, t in enumerate(page_texts, 1):
                        pages_data.append({
                            "page_number": i,
                            "text": t,
                            "char_count": len(t),
                            "ocr_applied": False,
                        })
                    logger.info("[EXTRACT] PyMuPDF succeeded: %d chars", len(full_text))
                else:
                    logger.info(
                        "[EXTRACT] PyMuPDF weak (%d chars) — falling back to raw read",
                        len(candidate),
                    )
                doc.close()
            except Exception as _e:
                logger.warning("[EXTRACT] PyMuPDF failed: %s", _e)

        # ── Strategy 3: Raw bytes (genuine text recovery only) ────────────
        # Only accepted if the recovered text has a reasonable alpha-char ratio
        # (>= 30%), which indicates real readable text rather than binary noise.
        # If the ratio check fails, full_text remains empty — no sentinel injected.
        if len(full_text) < 100:
            try:
                raw = file_path.read_bytes()
                decoded = raw.decode("latin-1", errors="replace")
                printable = "".join(
                    c if c.isprintable() or c in "\n\t" else " " for c in decoded
                )
                candidate = printable[:10000].strip()
                if candidate:
                    # Quality gate: reject if mostly binary/non-alpha content
                    alpha_ratio = sum(c.isalpha() for c in candidate) / len(candidate)
                    if alpha_ratio >= 0.30:
                        full_text = candidate
                        _extraction_method = "raw_bytes"
                        pages_data = [{
                            "page_number": 1,
                            "text": full_text,
                            "char_count": len(full_text),
                            "ocr_applied": False,
                        }]
                        logger.info(
                            "[EXTRACT] Raw bytes fallback accepted: %d chars (alpha_ratio=%.1f%%)",
                            len(full_text), alpha_ratio * 100,
                        )
                    else:
                        logger.warning(
                            "[EXTRACT] Raw bytes rejected: alpha_ratio=%.1f%% (< 30%%) — "
                            "treating as binary/unreadable",
                            alpha_ratio * 100,
                        )
            except Exception as _e:
                logger.error("[EXTRACT] Raw fallback failed: %s", _e)

        # ── No sentinel injected ─────────────────────────────────────────
        # If all three strategies failed, full_text remains "".
        # Failure is represented in metadata (extraction_failed / quality_tag),
        # NOT by injecting fake text. The downstream pipeline must handle empty
        # text via its own blocking logic.
        if not full_text.strip():
            logger.error(
                "[EXTRACT] ALL methods failed for %s — full_text is empty. "
                "Set extraction_failed=True in metadata.",
                file_path.name,
            )

        # ── Table extraction (preserves existing pipeline behaviour) ────
        # Only appended when we already have real extracted text.
        if full_text.strip():
            tables_text = self._extract_tables_from_pdf(file_path)
            if tables_text:
                if "SYNOPSIS TABLE" in tables_text:
                    synopsis_part = ""
                    other_part = ""
                    for block in tables_text.split("\n\n"):
                        if "SYNOPSIS TABLE" in block:
                            synopsis_part += block + "\n\n"
                        else:
                            other_part += block + "\n\n"
                    if synopsis_part:
                        full_text = (
                            f"\n--- Synopsis Content ---\n\n{synopsis_part.strip()}\n\n{full_text}"
                        )
                    if other_part.strip():
                        full_text += f"\n\n--- Extracted Tables ---\n\n{other_part.strip()}"
                else:
                    full_text += f"\n\n--- Extracted Tables ---\n\n{tables_text}"

        full_text = full_text.strip()

        # ── Extraction Confidence ──────────────────────────────────────
        ratio = len(full_text) / baseline_len
        if ratio > 0.7:
            extraction_confidence = "HIGH"
        elif ratio > 0.4:
            extraction_confidence = "MEDIUM"
        else:
            extraction_confidence = "LOW"
            
        self._last_extraction_confidence = extraction_confidence

        # ── Quality tag for downstream validation ───────────────────────
        if not full_text:
            quality_tag = "FAILED"
        elif len(full_text) < 200:
            quality_tag = "LOW_QUALITY"
            logger.warning(
                "[EXTRACT] LOW_QUALITY: %s extracted only %d chars via %s",
                file_path.name, len(full_text), _extraction_method,
            )
        else:
            quality_tag = "OK"

        # ── Mandatory [EXTRACT] debug trace ─────────────────────────────
        print(f"[EXTRACT] File: {file_path.name}")
        print(f"[EXTRACT] Method: {_extraction_method}")
        print(f"[EXTRACT] Length: {len(full_text)}")
        print(f"[EXTRACT] Quality: {quality_tag}")
        print(f"[EXTRACT] Preview: {full_text[:300]}")

        # Store extraction method on instance for load_document() to pick up
        self._last_extraction_method = _extraction_method
        self._last_extraction_quality = quality_tag

        return full_text, pages_data
    
    def _ocr_page(self, page) -> str:
        """
        Apply OCR to a PDF page using pytesseract.
        
        Args:
            page: A PyMuPDF page object
            
        Returns:
            OCR-extracted text
        """
        try:
            # Render page to image at 300 DPI
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)
            return text
        except Exception as e:
            logger.warning(f"OCR failed for page: {e}")
            return ""
    
    def _is_key_value_table(self, table: list) -> bool:
        """
        Detect if a table is a 2-column key-value layout (common in synopses).
        Key-value tables have exactly 2 columns where the first column contains
        short labels and the second column contains longer values.
        """
        if not table or len(table) < 2:
            return False
        col_counts = set(len(row) for row in table if row)
        if col_counts != {2}:
            return False
        # Check: first column cells should be shorter than second column cells
        label_lens = []
        value_lens = []
        for row in table:
            if row and len(row) == 2:
                c0 = str(row[0]).strip() if row[0] else ""
                c1 = str(row[1]).strip() if row[1] else ""
                if c0:
                    label_lens.append(len(c0))
                    value_lens.append(len(c1))
        if not label_lens:
            return False
        avg_label = sum(label_lens) / len(label_lens)
        avg_value = sum(value_lens) / len(value_lens) if value_lens else 0
        # Labels are typically shorter than values in key-value tables
        return avg_label < 80 and (avg_value > avg_label or avg_value > 20)

    def _format_key_value_table(self, table: list) -> str:
        """Format a 2-column key-value table as structured text."""
        lines = []
        for row in table:
            if not row or len(row) < 2:
                continue
            label = str(row[0]).strip() if row[0] else ""
            value = str(row[1]).strip() if row[1] else ""
            if label and value:
                # Multi-line values: normalize newlines
                value = ' '.join(value.split())
                lines.append(f"{label}: {value}")
            elif label:
                lines.append(f"{label}:")
            elif value:
                # Continuation of previous value
                if lines:
                    lines[-1] += f" {value}"
                else:
                    lines.append(value)
        return '\n'.join(lines)

    def _extract_tables_from_pdf(self, file_path: Path) -> str:
        """
        Extract tables from a PDF using pdfplumber.
        Detects key-value tables (common in synopses) and formats them as
        structured Label: Value text instead of flat pipe-delimited rows.
        
        Args:
            file_path: Path to the PDF
            
        Returns:
            Formatted table text
        """
        if not PDFPLUMBER_AVAILABLE:
            return ""
        
        try:
            tables_text = ""
            synopsis_text = ""
            with pdfplumber.open(str(file_path)) as pdf:
                table_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()
                    for table in tables:
                        if not table or len(table) < 1:
                            continue
                        table_count += 1
                        
                        # Check if this is a key-value table (synopsis-style)
                        if self._is_key_value_table(table):
                            kv_text = self._format_key_value_table(table)
                            # Check if this is likely a synopsis table
                            first_labels = ' '.join(
                                str(row[0]).strip().lower()
                                for row in table[:5]
                                if row and row[0]
                            )
                            if any(kw in first_labels for kw in [
                                'synopsis', 'title', 'protocol', 'study title',
                                'sponsor', 'phase', 'indication'
                            ]):
                                synopsis_text += f"\n\nSYNOPSIS TABLE (Page {page_num}):\n{kv_text}"
                            else:
                                tables_text += f"\nTable {table_count} (Page {page_num}):\n{kv_text}\n\n"
                        else:
                            # Standard table: format with pipe separators
                            tables_text += f"\n[TABLE_START]\n"
                            tables_text += f"Table {table_count} (Page {page_num}):\n"
                            for row in table:
                                cleaned = [str(cell).strip() if cell else "" for cell in row]
                                # Skip completely empty rows
                                if any(c for c in cleaned):
                                    tables_text += " | ".join(cleaned) + "\n"
                            tables_text += "[TABLE_END]\n\n"
            
            logger.info(f"Extracted {table_count} tables from {file_path.name}")
            # Synopsis tables go first (priority content)
            combined = ""
            if synopsis_text:
                combined += synopsis_text.strip() + "\n\n"
            if tables_text:
                combined += tables_text.strip()
            return combined.strip()
        except Exception as e:
            logger.warning(f"Table extraction failed for {file_path.name}: {e}")
            return ""
    
    def load_docx(self, file_path: Path) -> Tuple[str, List[Dict]]:
        """
        Load a DOCX file and extract text with structure.
        Extracts both paragraph text AND full table cell content to ensure
        structured clinical data (e.g., Synopsis tables) is captured.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (full_text, paragraphs_data)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")

        doc = Document(str(file_path))
        full_text = ""
        paragraphs_data = []

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                full_text += para.text + "\n"
                paragraphs_data.append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                })

        # ── Extract table content (including all cell text) ────────────
        # Tables in clinical docs often hold Synopsis / Demographics data
        # that is NOT captured by paragraph extraction alone.
        table_lines = []
        for table_idx, table in enumerate(doc.tables):
            table_lines.append(f"\n--- Table {table_idx + 1} ---")
            for row in table.rows:
                # Collect every non-empty cell, including multi-line cell text
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Normalise internal newlines within a cell
                        cell_text = " ".join(cell_text.split())
                        cells.append(cell_text)
                if cells:
                    table_lines.append(" | ".join(cells))

        if table_lines:
            full_text += "\n\n--- Extracted Tables ---\n" + "\n".join(table_lines)

        full_text = full_text.strip()

        # Quality tag
        if not full_text:
            quality_tag = "FAILED"
        elif len(full_text) >= 200:
            quality_tag = "OK"
        else:
            quality_tag = "LOW_QUALITY"

        # ── Mandatory [EXTRACT] debug output ─────────────────────────
        print(f"[EXTRACT] File: {file_path.name}")
        print(f"[EXTRACT] Method: docx")
        print(f"[EXTRACT] Length: {len(full_text)}")
        print(f"[EXTRACT] Quality: {quality_tag}")
        print(f"[EXTRACT] Preview: {full_text[:300]}")
        if not full_text:
            print(f"[EXTRACT] WARNING: Empty extraction for {file_path.name} — marked as unusable")

        self._last_extraction_method = "docx"
        self._last_extraction_quality = quality_tag
        self._last_extraction_confidence = "HIGH"
        return full_text, paragraphs_data
    
    def detect_document_type(self, text: str, filename: str) -> str:
        """
        Detect the type of clinical document based on content and filename.
        
        Args:
            text: The document text
            filename: The original filename
            
        Returns:
            Document type string
        """
        text_lower = text.lower()[:5000]  # Check first 5000 chars
        filename_lower = filename.lower()
        
        # Check filename first
        for doc_type, patterns in self.DOCUMENT_TYPE_PATTERNS.items():
            for pattern in patterns:
                if pattern in filename_lower:
                    return doc_type
                    
        # Check content
        for doc_type, patterns in self.DOCUMENT_TYPE_PATTERNS.items():
            for pattern in patterns:
                if pattern in text_lower:
                    return doc_type
                    
        return "unknown"
    
    def load_document(self, file_path: Path) -> Dict:
        """
        Load a single document and extract its content.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary with document metadata and content.
            Always includes 'extraction_method' and 'extraction_quality' fields.
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        # Reset per-file extraction state
        self._last_extraction_method = "none"
        self._last_extraction_quality = "unknown"
        self._last_extraction_confidence = "UNKNOWN"

        if extension == ".pdf":
            full_text, structure = self.load_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            full_text, structure = self.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

        doc_type = self.detect_document_type(full_text, file_path.name)

        # Extraction failed = text is genuinely empty (no sentinel, no fake text)
        extraction_failed = not bool(full_text.strip())
        if extraction_failed:
            logger.error(
                "[EXTRACT] FAIL — %s returned empty text after all extraction attempts.",
                file_path.name,
            )
            # Override quality to FAILED when text is truly empty
            self._last_extraction_quality = "FAILED"
            self._last_extraction_method = "none"

        document_data = {
            "filename": file_path.name,
            "filepath": str(file_path),
            "extension": extension,
            "document_type": doc_type,
            "full_text": full_text,
            "structure": structure,
            "char_count": len(full_text),
            "word_count": len(full_text.split()),
            "extraction_failed": extraction_failed,
            # Explicit extraction traceability fields
            "extraction_method": self._last_extraction_method,
            "extraction_quality": self._last_extraction_quality,
            "extraction_confidence": self._last_extraction_confidence,
        }

        self.loaded_documents[file_path.name] = document_data
        return document_data
    
    def validate_extraction(self, doc_data: Dict) -> Dict:
        """
        Validate extraction quality and flag potential issues.
        
        Provides 4-tier quality classification with per-page density metrics:
        - FAILED: <100 total chars (likely scanned/image PDF)
        - POOR: <30 chars/page (garbled or near-empty extraction)
        - WARNING: <100 chars/page or high special-char ratio
        - GOOD: ≥100 chars/page with clean text
        
        Args:
            doc_data: Document data dictionary
            
        Returns:
            Updated document data with extraction quality info and diagnostics
        """
        issues = []
        
        char_count = doc_data.get("char_count", 0)
        word_count = doc_data.get("word_count", 0)
        full_text = doc_data.get("full_text", "")
        structure = doc_data.get("structure", [])
        
        # Calculate page count from structure data
        page_count = len(structure) if structure else 1
        chars_per_page = char_count / max(page_count, 1)
        
        # Calculate alpha ratio (text quality indicator)
        alpha_ratio = 0.0
        if full_text:
            alpha_ratio = sum(c.isalpha() or c.isspace() for c in full_text) / max(len(full_text), 1)
        
        # --- Quality classification ---
        
        # Tier 1: FAILED - near-empty extraction
        if char_count < 100:
            issues.append("CRITICAL: Extraction returned almost no text - file may be scanned/image PDF without OCR")
        
        # Tier 2: POOR - very low density per page
        elif chars_per_page < 30:
            issues.append(f"CRITICAL: Very low text density ({chars_per_page:.0f} chars/page) - extraction likely failed for most pages")
        
        # Tier 3: WARNING - low density or quality issues
        elif chars_per_page < 100:
            issues.append(f"WARNING: Low text density ({chars_per_page:.0f} chars/page) - possible partial extraction")
        
        if char_count >= 100 and char_count < 500:
            issues.append("WARNING: Very short extraction - possible OCR issue or partial extraction")
        
        if word_count < 50 and char_count >= 100:
            issues.append("WARNING: Low word count - verify document is readable")
            
        # Check for garbled text (high ratio of special characters)
        if full_text and alpha_ratio < 0.5:
            issues.append(f"WARNING: High ratio of special characters ({alpha_ratio:.1%} alphanumeric) - extraction may be corrupted")
        
        # Determine quality level
        if any("CRITICAL" in i for i in issues):
            if char_count < 100 or chars_per_page < 30:
                quality = "failed" if char_count < 100 else "poor"
            else:
                quality = "poor"
        elif issues:
            quality = "warning"
        else:
            quality = "good"
            
        doc_data["extraction_quality"] = quality
        doc_data["extraction_issues"] = issues
        
        # Detect synopsis content
        synopsis_detected = False
        if full_text:
            text_lower = full_text.lower()
            synopsis_detected = any(kw in text_lower for kw in [
                'synopsis table', 'synopsis', 'study synopsis',
                'protocol synopsis', 'clinical study synopsis'
            ])
        
        # Structured diagnostics for downstream reporting
        doc_data["extraction_diagnostics"] = {
            "char_count": char_count,
            "word_count": word_count,
            "page_count": page_count,
            "chars_per_page": round(chars_per_page, 1),
            "alpha_ratio": round(alpha_ratio, 3),
            "quality_level": quality.upper(),
            "synopsis_detected": synopsis_detected,
            "issues": issues
        }
        
        return doc_data
    
    def load_all_documents(self) -> Dict[str, Dict]:
        """
        Load all documents from the input directory.
        
        Returns:
            Dictionary mapping filenames to document data
        """
        files = self.get_available_files()
        
        for file_path in files:
            try:
                doc_data = self.load_document(file_path)
                # Validate extraction quality
                doc_data = self.validate_extraction(doc_data)
                self.loaded_documents[file_path.name] = doc_data
                
                quality = doc_data.get("extraction_quality", "unknown")
                doc_type = doc_data.get("document_type", "unknown")
                
                if quality == "good":
                    print(f"✓ Loaded: {file_path.name} (Type: {doc_type})")
                elif quality == "warning":
                    print(f"⚠ Loaded with issues: {file_path.name} (Type: {doc_type})")
                    for issue in doc_data.get("extraction_issues", []):
                        print(f"    - {issue}")
                else:  # failed
                    print(f"✗ Extraction failed: {file_path.name}")
                    for issue in doc_data.get("extraction_issues", []):
                        print(f"    - {issue}")
            except Exception as e:
                print(f"✗ Failed to load {file_path.name}: {str(e)}")
                
        return self.loaded_documents
    
    def get_documents_by_type(self, doc_type: str) -> List[Dict]:
        """Get all loaded documents of a specific type."""
        return [
            doc for doc in self.loaded_documents.values()
            if doc["document_type"] == doc_type
        ]
    
    def save_extracted_text(self, output_dir: Path) -> None:
        """
        Save extracted text to JSON files for later use.
        
        Args:
            output_dir: Directory to save extracted text files
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for filename, doc_data in self.loaded_documents.items():
            # Create a serializable version (remove large structure data)
            save_data = {
                "filename": doc_data["filename"],
                "document_type": doc_data["document_type"],
                "full_text": doc_data["full_text"],
                "char_count": doc_data["char_count"],
                "word_count": doc_data["word_count"]
            }
            
            output_file = output_dir / f"{Path(filename).stem}_extracted.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, indent=2, ensure_ascii=False)
                
            print(f"✓ Saved: {output_file.name}")


# Convenience function
def load_documents_from_directory(input_dir: Path) -> Dict[str, Dict]:
    """Load all documents from a directory."""
    loader = DocumentLoader(input_dir)
    return loader.load_all_documents()
